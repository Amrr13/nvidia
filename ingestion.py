import os
import fitz  # PyMuPDF
import docx
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config import settings

def parse_pdf(file_path: str) -> str:
    """Extracts text from a PDF file."""
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def parse_docx(file_path: str) -> str:
    """Extracts text from a DOCX file."""
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def ingest_file(file_path: str) -> list[Document]:
    """
    Parses a document (PDF or DOCX), creates a LangChain Document,
    and chunks it into smaller pieces for vector storage.
    """
    if file_path.lower().endswith('.pdf'):
        text = parse_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        text = parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format for {file_path}")
    
    # Create a LangChain Document with metadata
    filename = os.path.basename(file_path)
    doc = Document(page_content=text, metadata={"source": filename})
    
    # Chunk the document
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP
    )
    chunks = text_splitter.split_documents([doc])
    
    # Update chunk metadata
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
        
    return chunks
